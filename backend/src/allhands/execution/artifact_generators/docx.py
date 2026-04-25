"""DOCX generator · structured block list → .docx bytes via python-docx.

Spec: docs/specs/2026-04-25-artifact-kinds-roadmap.md § 2.4.

Supported block types:
- heading      ``{"type": "heading", "level": 1-6, "text": str}``
- paragraph    ``{"type": "paragraph", "text": str}``
- list         ``{"type": "list", "ordered": bool, "items": list[str]}``
- code         ``{"type": "code", "language": str?, "text": str}``
- table        ``{"type": "table", "headers": list[str]?, "rows": list[list]}``

Unrecognised types are skipped with a warning so an agent that hallucinates
a block type doesn't blow up the whole document.
"""

from __future__ import annotations

import io
from typing import Any

from allhands.execution.artifact_generators.pdf import ArtifactGenerationError


def render_docx(*, blocks: list[Any]) -> tuple[bytes, list[str]]:
    """Returns ``(bytes, warnings)``. Warnings list non-fatal issues like
    ``"unknown block type: foo"`` so the executor can echo them in the
    tool result envelope.
    """
    if not isinstance(blocks, list):
        raise ArtifactGenerationError("blocks must be a list.")

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover
        raise ArtifactGenerationError(f"python-docx unavailable: {exc}") from exc

    warnings: list[str] = []
    doc = Document()
    # Default body font · Calibri stays generic Word-friendly.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for idx, block in enumerate(blocks):
        if not isinstance(block, dict):
            warnings.append(f"blocks[{idx}] is not an object — skipped")
            continue
        btype = block.get("type")
        try:
            if btype == "heading":
                level = int(block.get("level", 1))
                level = max(1, min(level, 9))  # python-docx caps at 9
                doc.add_heading(str(block.get("text", "")), level=level)
            elif btype == "paragraph":
                doc.add_paragraph(str(block.get("text", "")))
            elif btype == "list":
                ordered = bool(block.get("ordered", False))
                style_name = "List Number" if ordered else "List Bullet"
                items = block.get("items") or []
                for item in items:
                    doc.add_paragraph(str(item), style=style_name)
            elif btype == "code":
                p = doc.add_paragraph(str(block.get("text", "")))
                run = p.runs[0] if p.runs else p.add_run("")
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            elif btype == "table":
                headers = block.get("headers") or []
                rows = block.get("rows") or []
                cols = max(len(headers), max((len(r) for r in rows), default=0))
                if cols == 0:
                    warnings.append(f"blocks[{idx}] table has no columns — skipped")
                    continue
                tbl = doc.add_table(rows=1 if headers else 0, cols=cols)
                tbl.style = "Light Grid Accent 1"
                if headers:
                    hdr_row = tbl.rows[0].cells
                    for c, h in enumerate(headers):
                        if c < cols:
                            hdr_row[c].text = str(h)
                for row in rows:
                    cells = tbl.add_row().cells
                    for c, val in enumerate(row):
                        if c < cols:
                            cells[c].text = "" if val is None else str(val)
            else:
                warnings.append(f"blocks[{idx}] unknown type {btype!r} — skipped")
        except Exception as exc:
            warnings.append(f"blocks[{idx}] render failed: {exc}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), warnings
