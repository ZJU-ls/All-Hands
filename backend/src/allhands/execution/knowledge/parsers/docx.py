"""DOCX parser — uses ``python-docx`` if installed.

Walks paragraphs in document order; paragraphs with `Heading 1`..`Heading 6`
styles become Sections (char offsets pre-computed). Table cells are emitted
as plain rows separated by tabs — fine for FTS, agents can ask for the raw
file via `kb_read_document` when they need exact layout.

`python-docx` is an optional dep. Importing it here is wrapped in
try/except so missing-dep environments raise a clear `DocxParseError`
that the orchestrator turns into a clean FAILED state with remediation
hint, rather than a stack trace.
"""

from __future__ import annotations

import re

from allhands.execution.knowledge.parsers import ParseResult, Section


class DocxParseError(RuntimeError):
    pass


_HEADING_RE = re.compile(r"^Heading\s*([1-6])$", re.IGNORECASE)


class DocxParser:
    mime_types: tuple[str, ...] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def parse(self, file_path: str) -> ParseResult:
        try:
            import docx  # type: ignore[import-not-found]
        except ImportError as exc:
            raise DocxParseError(
                "python-docx is required to ingest DOCX. Install with `uv add python-docx`."
            ) from exc

        doc = docx.Document(file_path)
        parts: list[str] = []
        sections: list[Section] = []

        def _flush_paragraph(text: str, level: int | None) -> None:
            text = text.strip()
            if not text:
                # Preserve a blank line so chunker can break on it
                parts.append("\n")
                return
            offset = sum(len(p) for p in parts)
            if level is not None:
                marker = f"\n{text}\n\n"
                parts.append(marker)
                sections.append(
                    Section(
                        title=text,
                        level=level,
                        char_start=offset + 1,  # past leading \n
                        char_end=offset + 1 + len(text),  # placeholder; fixed-up below
                    )
                )
            else:
                parts.append(text + "\n\n")

        for para in doc.paragraphs:
            level = None
            style = getattr(para.style, "name", "") or ""
            m = _HEADING_RE.match(style)
            if m:
                level = int(m.group(1))
            _flush_paragraph(para.text, level)

        # Flatten tables → tab-separated rows after paragraphs (good enough
        # for v0; downstream chunker doesn't care).
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text + "\n")
            parts.append("\n")

        text = "".join(parts)

        # Fix up Section.char_end to span up to the next heading at same/higher level
        for i, sec in enumerate(sections):
            end = len(text)
            for j in range(i + 1, len(sections)):
                if sections[j].level <= sec.level:
                    end = sections[j].char_start
                    break
            sections[i] = sec.model_copy(update={"char_end": end})

        meta = {
            "format": "docx",
            "title": getattr(doc.core_properties, "title", None),
            "author": getattr(doc.core_properties, "author", None),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
        return ParseResult(text=text, sections=sections, metadata=meta)
